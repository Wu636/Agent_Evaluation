"""
题卷答案生成器
解析 Word 题卷，调用 LLM 生成5个不同等级的学生答案，并生成 .docx 文件
"""

import json
import os
import re
import asyncio
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
except ImportError:
    Document = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

import requests

# Import Cloud API functions
from homework_reviewer_v2 import upload_file, homework_file_analysis

# 模型名称映射：前端 id → API 实际需要的模型名
# 某些模型在 API 中需要特定格式的名称（如带空格的大写名）
MODEL_NAME_MAPPING = {
    "claude-sonnet-4.5": "Claude Sonnet 4.5",
    "claude-sonnet-4-6": "claude-sonnet-4-6",
    "claude-haiku-4.5": "Claude Haiku 4.5",
    "claude-opus-4": "Claude Opus 4",
    "claude-opus-4-6": "claude-opus-4-6",
    "gemini-2.5-pro": "gemini-2.5-pro",
    "gemini-2.5-flash": "gemini-2.5-flash",
    "grok-4": "grok-4",
}

TITLE_KEYWORDS = (
    "试卷", "题卷", "作业", "考试", "测试", "测验", "练习", "报告", "实验",
    "论文", "案例", "任务", "课程设计", "设计", "项目", "调研", "调查", "实践",
    "实训", "训练", "研究", "方案"
)
TITLE_BAD_PHRASES = (
    "具体要求如下", "要求如下", "完成一份", "请完成", "请结合", "请根据",
    "通过考察", "通过实验", "收集、获取", "答题要求", "作答要求"
)
TITLE_BAD_PREFIXES = ("请", "根据", "围绕", "阅读", "结合", "完成", "撰写", "说明", "分析", "以“", "以\"", "以'", "以‘")
DEFAULT_EXAM_TITLE = "作业"
GENERIC_TITLES = {"粘贴题卷", "题卷文本", "未命名题卷", "题卷内容", "题卷", DEFAULT_EXAM_TITLE}
MAX_FALLBACK_SOURCE_NAME_BYTES = 90
MAX_FALLBACK_SOURCE_NAME_CHARS = 40
MAX_OUTPUT_FILENAME_BYTES = 180


def _normalize_text(text: str) -> str:
    text = (text or "").replace("\u3000", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text.strip("-_:：·• ")


def _extract_non_empty_lines(text: str, limit: int = 8) -> List[str]:
    lines = []
    for line in text.splitlines():
        cleaned = _normalize_text(line)
        if cleaned:
            lines.append(cleaned)
        if len(lines) >= limit:
            break
    return lines


def _is_generic_title(title: str) -> bool:
    normalized = _normalize_text(title)
    return normalized in GENERIC_TITLES


def _is_likely_title(candidate: str) -> bool:
    text = _normalize_text(candidate)
    if not text or len(text) < 2:
        return False
    if len(text) > 60:
        return False
    if text.endswith(("。", "！", "？", "；")):
        return False
    if any(phrase in text for phrase in TITLE_BAD_PHRASES):
        return False
    if re.match(r"^(第?[0-9一二三四五六七八九十]+[、.．])", text):
        return False
    if re.match(r"^[（(]?[0-9一二三四五六七八九十]+[）)]", text):
        return False
    if any(text.startswith(prefix) for prefix in TITLE_BAD_PREFIXES) and len(text) > 10:
        return False

    sentence_punctuation = sum(text.count(ch) for ch in "，。！？；")
    has_keyword = any(keyword in text for keyword in TITLE_KEYWORDS)

    if sentence_punctuation >= 2:
        return False
    if "，" in text and not has_keyword:
        return False

    return has_keyword or sentence_punctuation == 0 and len(text) <= 25


def _pick_exam_title(full_text: str, fallback_title: str = "", preferred_title: str = "") -> str:
    candidates = []
    preferred = _normalize_text(preferred_title)
    if preferred and not _is_generic_title(preferred):
        candidates.append(preferred)
    candidates.extend(_extract_non_empty_lines(full_text, limit=8))

    for candidate in candidates:
        if _is_likely_title(candidate):
            return candidate

    fallback = _normalize_text(fallback_title or preferred_title or DEFAULT_EXAM_TITLE)
    return fallback or DEFAULT_EXAM_TITLE


def _fallback_title_from_filename(file_stem: str) -> str:
    normalized = _normalize_text(file_stem)
    if not normalized:
        return DEFAULT_EXAM_TITLE
    if _is_generic_title(normalized):
        return DEFAULT_EXAM_TITLE
    if (
        len(normalized) > MAX_FALLBACK_SOURCE_NAME_CHARS
        or len(normalized.encode("utf-8")) > MAX_FALLBACK_SOURCE_NAME_BYTES
    ):
        return DEFAULT_EXAM_TITLE
    return normalized


def _truncate_utf8(text: str, max_bytes: int) -> str:
    if len(text.encode("utf-8")) <= max_bytes:
        return text

    current_bytes = 0
    result = []
    for ch in text:
        ch_bytes = len(ch.encode("utf-8"))
        if current_bytes + ch_bytes > max_bytes:
            break
        result.append(ch)
        current_bytes += ch_bytes
    return "".join(result).rstrip("._- ")


def _sanitize_filename_component(text: str, fallback: str = DEFAULT_EXAM_TITLE) -> str:
    normalized = _normalize_text(text)
    sanitized = re.sub(r'[\\/:*?"<>|]', "_", normalized)
    sanitized = re.sub(r"\s+", "_", sanitized).strip("._- ")
    return sanitized or fallback


def build_output_filename(title: str, file_suffix: str, fallback_title: str = DEFAULT_EXAM_TITLE) -> str:
    safe_title = _sanitize_filename_component(title, fallback=fallback_title)
    suffix = _sanitize_filename_component(file_suffix, fallback="学生答案")
    title_budget = max(
        24,
        MAX_OUTPUT_FILENAME_BYTES - len(suffix.encode("utf-8")) - len(".docx".encode("utf-8")) - 1,
    )
    safe_title = _truncate_utf8(safe_title, title_budget)
    if not safe_title:
        safe_title = fallback_title
    return f"{safe_title}_{suffix}.docx"


def load_llm_config_from_args(context: dict) -> Tuple[str, str, str]:
    """从上下文加载 LLM 配置"""
    api_key = context.get("llm_api_key") or os.getenv("LLM_API_KEY", "")
    api_url = context.get("llm_api_url") or os.getenv("LLM_API_URL", "http://llm-service.polymas.com/api/openai/v1/chat/completions")
    raw_model = context.get("llm_model") or os.getenv("LLM_MODEL", "claude-sonnet-4-6")
    # 应用模型名映射（前端可能传了 id 如 "claude-sonnet-4.5"，API 需要 "Claude Sonnet 4.5"）
    model = MODEL_NAME_MAPPING.get(raw_model, raw_model)
    return api_key, api_url, model


def extract_questions_from_cloud(docx_path: Path, context: dict) -> Tuple[str, str]:
    """
    使用云端 API 解析题卷结构：
    1. upload_file  → 获取 fileUrl
    2. homework_file_analysis → 获取 textInput 结构
    """
    authorization = os.getenv("AUTHORIZATION", "")
    cookie_env = os.getenv("COOKIE", "")
    if not authorization or not cookie_env:
        print(f"📄 使用本地解析模式")
        return extract_questions_from_local(docx_path)

    auth_preview = authorization[:20] + "..." if len(authorization) > 20 else authorization
    print(f"☁️ 正在上传题卷到云端: {docx_path.name} (auth={auth_preview})")
    try:
        # Step 1: 上传文件获取 fileUrl
        file_info = upload_file(str(docx_path))
        if not file_info or not file_info.get("fileUrl"):
            raise ValueError("文件上传失败，未获取到 fileUrl。请检查 Authorization/Cookie 是否已过期")

        print(f"✅ 文件上传成功: {file_info.get('fileName')}")

        # Step 2: 调用 homework_file_analysis 解析题卷结构
        print(f"☁️ 正在解析题卷结构...")
        success, result, text_input = homework_file_analysis(file_info, context)

        if not success or not text_input:
            error_msg = ""
            if isinstance(result, dict):
                error_msg = result.get("error", "") or result.get("msg", "")
            raise ValueError(f"云端解析失败: {error_msg}")

        print(f"✅ 题卷结构解析完成")

        parsed = text_input
        if isinstance(text_input, str):
            try:
                parsed = json.loads(text_input)
            except json.JSONDecodeError:
                # 纯文本直接用
                title = _pick_exam_title(text_input, fallback_title=_fallback_title_from_filename(docx_path.stem))
                return title, text_input

        # 如果是列表结构 [{ itemName, stuAnswerContent }, ...]
        # 格式化为可读文本
        if isinstance(parsed, list):
            full_text_lines = []
            for i, item in enumerate(parsed):
                item_name = item.get("itemName", "")
                content = item.get("stuAnswerContent", "")
                if item_name or content:
                    full_text_lines.append(f"【{item_name}】" if item_name else f"题目{i + 1}:")
                    if content:
                        full_text_lines.append(content)
                    full_text_lines.append("")
            full_text = "\n".join(full_text_lines)
            title = _pick_exam_title(full_text, fallback_title=_fallback_title_from_filename(docx_path.stem))
            return title, full_text

        if isinstance(parsed, dict):
            serialized = json.dumps(parsed, ensure_ascii=False, indent=2)
            title = _pick_exam_title(serialized, fallback_title=_fallback_title_from_filename(docx_path.stem))
            return title, serialized

        raw_text = str(parsed)
        title = _pick_exam_title(raw_text, fallback_title=_fallback_title_from_filename(docx_path.stem))
        return title, raw_text

    except Exception as e:
        print(f"⚠️ 云端解析失败 ({e})，尝试本地解析...")
        return extract_questions_from_local(docx_path)


def extract_questions_from_local(file_path: Path) -> Tuple[str, str]:
    """
    本地解析作为兜底，支持 .docx / .doc / .pdf
    """
    ext = file_path.suffix.lower()
    
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext in (".doc", ".docx"):
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"本地解析不支持 {ext} 格式，仅支持 .docx / .doc / .pdf。请配置智慧树认证信息以使用云端解析。")


def extract_questions_from_text(source_text: str, preferred_title: str = "") -> Tuple[str, str]:
    """解析用户粘贴的题卷文本"""
    exam_content = source_text.strip()
    if not exam_content:
        raise ValueError("题卷文字内容为空")
    fallback_title = DEFAULT_EXAM_TITLE
    title = _pick_exam_title(exam_content, fallback_title=fallback_title, preferred_title=preferred_title)
    return title, exam_content


def _extract_from_docx(docx_path: Path) -> Tuple[str, str]:
    """解析 .docx 文件"""
    if Document is None:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    doc = Document(docx_path)

    full_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text.strip())

    full_text_str = "\n".join(full_text)
    title = _pick_exam_title(full_text_str, fallback_title=_fallback_title_from_filename(docx_path.stem))
    return title, full_text_str


def _extract_from_pdf(pdf_path: Path) -> Tuple[str, str]:
    """解析 PDF 文件（使用 PyMuPDF）"""
    if fitz is None:
        raise ImportError("解析 PDF 需要 PyMuPDF，请安装: pip install PyMuPDF")
    
    doc = fitz.open(str(pdf_path))
    page_count = len(doc)
    
    full_text = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            full_text.append(text.strip())
    doc.close()
    
    all_text = "\n".join(full_text)
    
    title = _pick_exam_title(all_text, fallback_title=_fallback_title_from_filename(pdf_path.stem))
    
    print(f"✅ PDF 解析完成，共 {page_count} 页，提取 {len(all_text)} 字符")
    return title, all_text


def build_generation_prompt(title: str, exam_content: str, level: str, level_desc: str, custom_template: str = "") -> str:
    """构建生成答案的 Prompt"""
    if custom_template.strip():
        # 用户自定义模板：替换占位符
        return (
            custom_template
            .replace("{{title}}", title)
            .replace("{{level}}", level)
            .replace("{{level_desc}}", level_desc)
            .replace("{{exam_content}}", exam_content[:15000])
        )
    return f"""你是一名【{level}】水平的学生，正在作答《{title}》。
请根据你的水平要求完成所有题目或任务。

【等级要求：{level}】
{level_desc}

【作业内容】
{exam_content[:15000]}

【输出要求】
1. 请自动识别作业类型（试卷、论文、报告、案例分析、实验报告等），并按照对应格式作答
2. 如果是试卷类作业：按题型分类作答，保持题号对应，选择题紧凑排列
3. 如果是论文/报告类作业：输出完整的、符合题意和字数要求的文章
4. 如果是案例分析类作业：结合案例进行分析论述
5. 不要包含任何多余的开场白、解释或元评论，直接输出答案内容
6. 答案的质量必须严格符合【{level}】水平的设定
7. 如果是较低等级，应体现出知识理解不深入、存在错误或遗漏等特征
8. 绝对禁止使用任何 LaTeX 或 Markdown 数学语法，所有数学公式必须用纯文本表示。例如写 A^(-1) 而不是 $A^{{-1}}$，写 x=2 而不是 $x=2$
"""


async def generate_answer_content(prompt: str, context: dict) -> Optional[str]:
    """调用 LLM 生成答案内容"""
    api_key, api_url, model = load_llm_config_from_args(context)
    
    if not api_key:
        print("❌ 未配置 LLM API Key，请在右上角 ⚙️ 设置中配置")
        print(f"   context keys: {list(context.keys())}")
        print(f"   env LLM_API_KEY set: {bool(os.getenv('LLM_API_KEY'))}")
        return None
    
    # 首次调用时打印配置（脱敏）
    key_preview = api_key[:8] + "..." + api_key[-4:] if len(api_key) > 12 else "***"
    print(f"🔧 LLM 配置: url={api_url}, model={model}, key={key_preview}")
        
    headers = {
        "api-key": api_key,
        "Content-Type": "application/json"
    }
    
    payload = {
        "maxTokens": 4096,
        "messages": [
            {
                "role": "system",
                "content": "你是一个模拟真实学生作答的AI助手。你会根据指定的能力等级，生成符合该水平的作业答案。重要规则：所有输出必须是纯文本格式，绝对禁止使用 LaTeX/Markdown 数学语法（如 $...$、\\begin{}、\\frac{} 等），因为输出将写入 Word 文档。"
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "model": model, 
        "temperature": 0.5,
    }
    
    # 构建流式URL：普通URL + /stream 后缀（避免代理服务器120秒网关超时）
    stream_url = api_url.rstrip("/") + "/stream" if not api_url.endswith("/stream") else api_url
    
    max_retries = 3
    for attempt in range(max_retries + 1):
        try:
            # 使用流式接口，逐块读取响应，避免 Nginx 504 Gateway Timeout
            loop = asyncio.get_event_loop()
            
            def stream_request():
                """流式请求LLM API，拼接完整内容返回"""
                resp = requests.post(stream_url, headers=headers, json=payload, timeout=300, stream=True)
                if resp.status_code != 200:
                    body_preview = resp.text[:500] if resp.text else "(empty)"
                    raise requests.exceptions.HTTPError(
                        f"LLM API 返回 {resp.status_code}: {body_preview}",
                        response=resp
                    )
                
                # 从SSE流中拼接完整内容
                # 兼容两种格式：
                #   标准 OpenAI: "data: {...}" + choices[0].delta.content
                #   代理服务器: "data:{...}"  + choices[0].message.content
                full_content = []
                for line in resp.iter_lines(decode_unicode=True):
                    if not line:
                        continue
                    # 兼容 "data: {...}" 和 "data:{...}" 两种格式
                    if line.startswith("data:"):
                        data_str = line[5:].strip()
                    else:
                        continue
                    if data_str == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data_str)
                        choices = chunk.get("choices", [])
                        if choices:
                            choice = choices[0]
                            # 兼容 delta.content（标准）和 message.content（代理）
                            delta = choice.get("delta") or choice.get("message") or {}
                            content = delta.get("content")
                            if content:  # 跳过 null 和空字符串
                                full_content.append(content)
                    except json.JSONDecodeError:
                        continue
                return "".join(full_content)
            
            result = await loop.run_in_executor(None, stream_request)
            if result:
                return result
            
            print(f"⚠️ LLM 流式响应内容为空")
            return None
            
        except requests.exceptions.HTTPError as e:
            print(f"❌ LLM API 错误: {e}")
            print(f"   请求 URL: {stream_url}")
            print(f"   请求 Model: {model}")
            print(f"   API Key 前缀: {api_key[:10]}..." if len(api_key) > 10 else f"   API Key: (len={len(api_key)})")
            if attempt < max_retries:
                wait = 5 * (attempt + 1)
                print(f"   ⏳ 第 {attempt + 1} 次重试（等待 {wait}s）...")
                await asyncio.sleep(wait)
                continue
            return None
        except Exception as e:
            print(f"❌ LLM 生成失败: {type(e).__name__}: {e}")
            if attempt < max_retries:
                wait = 5 * (attempt + 1)
                print(f"   ⏳ 第 {attempt + 1} 次重试（等待 {wait}s）...")
                await asyncio.sleep(wait)
                continue
            return None


def create_answer_docx(content: str, output_path: Path, title: str, level: str, level_desc: str):
    """将生成的文本写入 Word 文档，模仿标准格式"""
    doc = Document()
    
    # 1. 试卷标题
    p_title = doc.add_paragraph(f"{title}五等级学生答案")
    if p_title.runs: p_title.runs[0].bold = True
    
    # 2. 等级描述
    p_level = doc.add_paragraph(f"等级：{level}（{level_desc}）")
    if p_level.runs: p_level.runs[0].bold = True 
    
    # 3. 写入内容
    # 简单处理：按行写入，识别到题型标题加粗
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        p = doc.add_paragraph(line)
        
        # 识别题型标题加粗 (一、xxx)
        if re.match(r'^[一二三四五六七八九十]+[、\.]', line):
            if p.runs: p.runs[0].bold = True
            
    doc.save(output_path)


LEVEL_DEFINITIONS = {
    "优秀的回答": "知识全面精准，逻辑清晰连贯，案例结合到位，合规细节无遗漏",
    "良好的回答": "覆盖较全，逻辑较清晰，有一定案例结合，偶有小瑕疵",
    "中等的回答": "基本知识点掌握，逻辑一般，案例结合较少，表述平铺直叙",
    "合格的回答": "核心知识点有遗漏，逻辑不够严密，表述存在模糊之处",
    "较差的回答": "知识漏洞多，逻辑混乱，未结合案例，存在明显错误"
}

LEVEL_FILENAMES = {
    "优秀的回答": "等级一_优秀_学生答案",
    "良好的回答": "等级二_良好_学生答案",
    "中等的回答": "等级三_中等_学生答案",
    "合格的回答": "等级四_合格_学生答案",
    "较差的回答": "等级五_较差_学生答案"
}


async def generate_level_answers(
    exam_docx_path: Optional[Path],
    output_dir: Path,
    levels: List[str],  # e.g., ["优秀的回答", "较差的回答"]
    context: dict,
    custom_prompt: str = "",
    custom_levels_json: str = "",
    source_text: str = "",
    source_title: str = "",
) -> List[Path]:
    """
    生成指定等级的答案文件
    """
    try:
        if source_text.strip():
            print("📝 正在解析粘贴题卷文字...")
            title, exam_content = extract_questions_from_text(source_text, preferred_title=source_title)
        elif exam_docx_path is not None:
            print(f"📄 正在解析题卷: {exam_docx_path.name}")
            title, exam_content = extract_questions_from_cloud(exam_docx_path, context)
        else:
            raise ValueError("缺少题卷输入内容")
    except Exception as e:
        print(f"❌ 解析题卷失败: {e}")
        return []

    print(f"✅ 题卷解析完成，标题: {title}，目标生成 {len(levels)} 份答案")
    
    # 解析自定义等级描述
    effective_level_defs = dict(LEVEL_DEFINITIONS)  # 先拷贝默认值
    if custom_levels_json:
        try:
            custom_defs = json.loads(custom_levels_json)
            if isinstance(custom_defs, dict):
                effective_level_defs.update(custom_defs)
                print(f"📝 已加载用户自定义等级描述")
        except json.JSONDecodeError:
            print(f"⚠️ 自定义等级描述 JSON 解析失败，使用默认值")
    
    output_dir.mkdir(parents=True, exist_ok=True)
    generated_files = []
    
    # 并发生成
    tasks = []
    for level_key in levels:
        # 映射 level key 到描述
        # 前端传来的可能是 "优秀", "良好" 等简写，需要匹配
        full_key = next((k for k in LEVEL_DEFINITIONS if level_key in k), level_key)
        
        desc = effective_level_defs.get(full_key, LEVEL_DEFINITIONS.get(full_key, "无描述"))
        file_suffix = LEVEL_FILENAMES.get(full_key, f"{level_key}_学生答案")

        filename = build_output_filename(
            title=title,
            file_suffix=file_suffix,
            fallback_title=_fallback_title_from_filename(exam_docx_path.stem) if exam_docx_path is not None else DEFAULT_EXAM_TITLE,
        )
        output_path = output_dir / filename
        
        tasks.append((full_key, desc, output_path))

    # 执行生成任务
    for level, desc, path in tasks:
        print(f"🤖 正在生成: {level}...")
        prompt = build_generation_prompt(title, exam_content, level, desc, custom_template=custom_prompt)
        content = await generate_answer_content(prompt, context)
        
        if content:
            create_answer_docx(content, path, title, level, desc)
            print(f"✅ 生成完毕: {path.name}")
            generated_files.append(path)
        else:
            print(f"❌ 生成失败: {level}")

    return generated_files
