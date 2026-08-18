"""使用 AgentEval 大模型生成作业批阅 Skill ZIP 与学生 DOCX 测试样例。"""

from __future__ import annotations

import base64
import json
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

try:
    from .skill_review_service import CorrectionSkillError, validate_grading_skill_package
except ImportError:
    from skill_review_service import CorrectionSkillError, validate_grading_skill_package


DEFAULT_LLM_API_URL = "https://llm-service.polymas.com/api/openai/v1/chat/completions"
MAX_MATERIAL_CHARS_PER_FILE = 18_000
MAX_MATERIAL_CONTEXT_CHARS = 60_000
SKILL_NAME_PATTERN = re.compile(r"^[a-z0-9]+(?:-[a-z0-9]+)*$")
SUPPORTED_MATERIAL_SUFFIXES = {".txt", ".md", ".csv", ".docx", ".pdf", ".xlsx"}

README_TEMPLATE = """# 批阅技能 · 文件结构说明

每个批阅技能 = 一种作业类型的完整批阅能力，按以下结构组织。

```text
<skill-name>/
├── README.md
├── SKILL.md
├── references/       # 可选；有实际参考文件时保留
└── scripts/          # 可选；有确定性脚本时保留
```

`SKILL.md` 定义批阅对象、取证流程、评分项和评价项；`references/` 只保存执行时需要读取的课程规则与匿名校准说明。
"""

SAMPLE_LEVELS = [
    ("优秀", "内容完整、证据充分、逻辑清楚，只有极少量轻微瑕疵"),
    ("良好", "核心内容完整，论证较清楚，但细节或深度略有不足"),
    ("中等", "完成主要任务，存在若干遗漏、表述一般或分析不够深入"),
    ("合格", "达到基本提交要求，但核心内容有明显缺漏或论证较弱"),
    ("较差", "仅完成少量要求，存在较多遗漏、错误或结构问题"),
]

SAMPLE_LABEL_LEAK_PATTERNS = (
    "优秀作业",
    "良好作业",
    "中等作业",
    "合格作业",
    "较差作业",
    "优秀档位",
    "良好档位",
    "中等档位",
    "合格档位",
    "较差档位",
    "内部目标档位",
    "AI生成",
    "AI 生成",
    "测试样例",
    "故意写差",
)


class SkillGenerationError(RuntimeError):
    """AgentEval LLM 生成或产物组装失败。"""


def normalize_chat_completion_endpoint(api_url: str) -> str:
    endpoint = (api_url or DEFAULT_LLM_API_URL).strip().rstrip("/")
    if not endpoint:
        return DEFAULT_LLM_API_URL
    if "/chat/completions" in endpoint:
        return endpoint
    return f"{endpoint}/chat/completions"


def _extract_llm_content(payload: Dict[str, Any]) -> str:
    choices = payload.get("choices")
    if not isinstance(choices, list) or not choices:
        raise SkillGenerationError("AgentEval 大模型响应缺少 choices")
    message = choices[0].get("message") if isinstance(choices[0], dict) else None
    content = message.get("content") if isinstance(message, dict) else None
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        text_parts = []
        for item in content:
            if isinstance(item, dict):
                value = item.get("text") or item.get("content")
                if isinstance(value, str):
                    text_parts.append(value)
        if text_parts:
            return "".join(text_parts).strip()
    raise SkillGenerationError("AgentEval 大模型响应内容为空")


def call_agenteval_llm(
    *,
    system_prompt: str,
    user_prompt: str,
    api_key: str,
    api_url: str,
    model: str,
    max_tokens: int = 12_000,
    timeout_seconds: int = 600,
) -> str:
    """调用 AgentEval 全局设置对应的 OpenAI 兼容接口。"""
    import requests

    endpoint = normalize_chat_completion_endpoint(api_url)
    normalized_model = model.strip()
    if not api_key.strip():
        raise SkillGenerationError("请先在 AgentEval 设置中配置 LLM API Key")
    if not normalized_model:
        raise SkillGenerationError("请先在 AgentEval 设置中配置作业批阅模型")

    headers = {
        "api-key": api_key.strip(),
        "Authorization": f"Bearer {api_key.strip()}",
        "Content-Type": "application/json",
    }
    request_payload = {
        "model": normalized_model,
        "maxTokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "n": 1,
    }
    try:
        response = requests.post(
            endpoint,
            headers=headers,
            json=request_payload,
            timeout=timeout_seconds,
        )
    except requests.RequestException as exc:
        raise SkillGenerationError(f"连接 AgentEval 大模型失败：{exc}") from exc

    if not response.ok:
        preview = response.text[:1000]
        raise SkillGenerationError(
            f"AgentEval 大模型请求失败（HTTP {response.status_code}）：{preview}"
        )
    try:
        payload = response.json()
    except ValueError as exc:
        raise SkillGenerationError("AgentEval 大模型返回了非 JSON 响应") from exc
    return _extract_llm_content(payload)


def parse_json_object(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    fenced = re.search(r"```(?:json)?\s*([\s\S]*?)```", cleaned, re.I)
    if fenced:
        cleaned = fenced.group(1).strip()
    try:
        value = json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start < 0 or end <= start:
            raise SkillGenerationError("大模型未返回 JSON 对象")
        try:
            value = json.loads(cleaned[start:end + 1])
        except json.JSONDecodeError as exc:
            raise SkillGenerationError(f"大模型 JSON 解析失败：{exc}") from exc
    if not isinstance(value, dict):
        raise SkillGenerationError("大模型结果必须是 JSON 对象")
    return value


def _truncate_text(text: str, limit: int) -> str:
    normalized = text.replace("\x00", "").strip()
    return normalized if len(normalized) <= limit else normalized[:limit] + "\n[内容已截断]"


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"[表格 {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("\t".join(cells))
    return "\n".join(lines)


def _extract_pdf_text(path: Path) -> str:
    import fitz

    document = fitz.open(str(path))
    try:
        return "\n".join(page.get_text("text") for page in document)
    finally:
        document.close()


def _extract_xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: List[str] = []
    try:
        for sheet in workbook.worksheets:
            lines.append(f"[工作表：{sheet.title}]")
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    lines.append("\t".join(values))
                if row_index >= 300:
                    lines.append("[工作表内容已截断]")
                    break
    finally:
        workbook.close()
    return "\n".join(lines)


def extract_material_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_MATERIAL_SUFFIXES:
        raise SkillGenerationError(f"课程材料暂不支持 {suffix or '未知'} 格式")
    if suffix in {".txt", ".md", ".csv"}:
        text = path.read_text(encoding="utf-8-sig", errors="replace")
    elif suffix == ".docx":
        text = _extract_docx_text(path)
    elif suffix == ".pdf":
        text = _extract_pdf_text(path)
    else:
        text = _extract_xlsx_text(path)
    return _truncate_text(text, MAX_MATERIAL_CHARS_PER_FILE)


def build_material_context(material_paths: Sequence[Path], material_text: str) -> str:
    sections: List[str] = []
    if material_text.strip():
        sections.append("【教师补充说明】\n" + _truncate_text(material_text, 20_000))
    for index, path in enumerate(material_paths, start=1):
        extracted = extract_material_text(path)
        sections.append(
            f"【材料 {index}｜文件名：{path.name}｜格式：{path.suffix.lower()}】\n{extracted}"
        )
    context = "\n\n".join(sections).strip()
    if not context:
        raise SkillGenerationError("请上传课程材料或填写教师补充说明")
    return _truncate_text(context, MAX_MATERIAL_CONTEXT_CHARS)


def _string_list(value: Any) -> List[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def validate_skill_blueprint(blueprint: Dict[str, Any]) -> List[str]:
    errors: List[str] = []
    skill_name = str(blueprint.get("skillName") or "").strip()
    if not SKILL_NAME_PATTERN.fullmatch(skill_name) or len(skill_name) > 64:
        errors.append("skillName 必须为不超过64字符的小写英文连字符名称")
    if len(str(blueprint.get("displayName") or "").strip()) < 4:
        errors.append("displayName 需要是清晰的教师端中文批阅名称")
    if len(str(blueprint.get("description") or "").strip()) < 30:
        errors.append("description 过短")
    if len(str(blueprint.get("submissionRequirement") or "").strip()) < 40:
        errors.append("submissionRequirement 过短")

    score_type = str(blueprint.get("scoreType") or "").strip().lower()
    if score_type not in {"dimension", "item"}:
        errors.append("scoreType 仅支持 dimension 或 item")
    try:
        full_score = int(blueprint.get("fullScore"))
    except (TypeError, ValueError):
        full_score = 0
    if full_score <= 0:
        errors.append("fullScore 必须为正整数")

    score_items = blueprint.get("scoreItems")
    if not isinstance(score_items, list) or not score_items:
        errors.append("scoreItems 至少包含一项")
    else:
        total = 0
        for index, item in enumerate(score_items, start=1):
            if not isinstance(item, dict):
                errors.append(f"scoreItems[{index}] 不是对象")
                continue
            if not str(item.get("name") or "").strip():
                errors.append(f"scoreItems[{index}] 缺少 name")
            try:
                score = int(item.get("score"))
            except (TypeError, ValueError):
                score = 0
            if score <= 0:
                errors.append(f"scoreItems[{index}] score 必须为正整数")
            total += score
            if len(str(item.get("description") or "").strip()) < 8:
                errors.append(f"scoreItems[{index}] description 过短")
            if not _string_list(item.get("rules")):
                errors.append(f"scoreItems[{index}] 缺少 rules")
        if full_score and total != full_score:
            errors.append(f"scoreItems 分值合计 {total}，应为 {full_score}")

    if len(_string_list(blueprint.get("workflow"))) < 3:
        errors.append("workflow 至少需要3步")
    if not _string_list(blueprint.get("missingRules")):
        errors.append("missingRules 不能为空")
    evaluation_names = {
        str(item.get("name") or "").strip()
        for item in blueprint.get("evaluationItems", [])
        if isinstance(item, dict)
    }
    if not {"综合评语", "改进建议"}.issubset(evaluation_names):
        errors.append("evaluationItems 必须包含综合评语和改进建议")
    return errors


def build_skill_blueprint_prompt(material_context: str, validation_feedback: str = "") -> str:
    feedback_section = (
        f"\n\n【上一次结果的校验问题】\n{validation_feedback}\n请修正全部问题。"
        if validation_feedback
        else ""
    )
    return f"""请把以下课程材料转换为一个可直接执行的作业批阅 Skill 蓝图。

必须遵守：
1. 区分权威题目、评分标准、教师补充说明、教师已评分样本和普通样本；冲突时以教师最新说明和最终版文件优先。
2. 教师样本只用于匿名校准，不得输出学生姓名、学号、原文、原始文件名或从文件名标签直接赋分的规则。
3. 学生提交要求必须写清“做什么、交什么、怎么交、关键质量要求”，不得出现历史样本校准、模型内部流程等术语。
4. 评分优先选择 dimension；固定题目或固定成果逐项评分时选择 item。总分与所有 scoreItems 分值之和必须完全一致。
5. 缺项、等价格式、重复扣分、材料读取受限和人工复核规则必须明确。
6. skillName 使用小写英文字母、数字和连字符，不超过64字符；description 至少30个汉字并说明触发场景。
7. 只输出一个合法 JSON 对象，不要 Markdown 代码块或解释。

JSON 结构：
{{
  "skillName": "english-kebab-name",
  "displayName": "8至20字中文名称",
  "description": "批阅对象、参考标准、评分评价与触发范围",
  "submissionRequirement": "可直接给学生看的完整要求",
  "scoreType": "dimension或item",
  "fullScore": 100,
  "itemSplit": "item类型时说明如何拆项，dimension可留空",
  "scoreItems": [
    {{"key": "ascii-key", "name": "评分项中文名", "score": 20, "description": "评分内容", "rules": ["可核查规则1", "可核查规则2"]}}
  ],
  "workflow": ["完整读取所有附件并建立证据清单", "逐项评分", "输出前审计"],
  "evidenceRules": ["证据读取和格式等价规则"],
  "missingRules": ["缺项、封顶、单一归因规则"],
  "courseRules": ["从材料提炼的匿名课程硬规则"],
  "calibrationNotes": ["仅保留匿名质量特征；无教师样本时写暂无样本校准"],
  "evaluationItems": [
    {{"name": "综合评语", "description": "整体表现概括"}},
    {{"name": "改进建议", "description": "下一次提交的可执行建议"}}
  ]
}}

【课程材料】
{material_context}{feedback_section}"""


def generate_skill_blueprint(
    *,
    material_context: str,
    api_key: str,
    api_url: str,
    model: str,
    max_attempts: int = 2,
) -> Dict[str, Any]:
    feedback = ""
    last_error = ""
    for _ in range(max(1, max_attempts)):
        response_text = call_agenteval_llm(
            system_prompt=(
                "你是作业批阅 Skill 架构师。严格依据用户材料生成可复核评分蓝图，"
                "保护学生隐私，输出必须是单一 JSON 对象。"
            ),
            user_prompt=build_skill_blueprint_prompt(material_context, feedback),
            api_key=api_key,
            api_url=api_url,
            model=model,
        )
        try:
            blueprint = parse_json_object(response_text)
            errors = validate_skill_blueprint(blueprint)
        except SkillGenerationError as exc:
            errors = [str(exc)]
            blueprint = {}
        if not errors:
            return blueprint
        last_error = "；".join(errors)
        feedback = last_error
    raise SkillGenerationError(f"大模型生成的 Skill 蓝图未通过校验：{last_error}")


def _escape_table_cell(value: Any) -> str:
    return str(value or "").strip().replace("|", "\\|").replace("\n", "<br>")


def _render_numbered_steps(items: Iterable[str]) -> str:
    return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1))


def render_skill_markdown(blueprint: Dict[str, Any]) -> str:
    name = str(blueprint["skillName"]).strip()
    display_name = str(blueprint["displayName"]).strip()
    description = re.sub(r"\s+", " ", str(blueprint["description"])).strip()
    score_type = str(blueprint["scoreType"]).strip().lower()
    full_score = int(blueprint["fullScore"])
    items = blueprint["scoreItems"]

    lines = [
        "---",
        f"name: {name}",
        f"description: {json.dumps(description, ensure_ascii=False)}",
        "---",
        "",
        f"# {display_name}",
        "",
        "## 批阅对象",
        "",
        "批阅学生提交的课程作业附件。支持以 DOCX、PDF、图片、表格或题目明确允许的等价格式提交；先完整读取可访问的正文、表格与插图，再依据 [课程规则与校准说明](references/course-rules.md) 评分。",
        "",
        "姓名、学号、班级、文件名中的分数或“优秀/中等/较差”等标签不得作为评分依据。材料读取不完整时列入人工复核，不把未确认事实写成错误。",
        "",
        "## 批阅流程",
        "",
        _render_numbered_steps(_string_list(blueprint.get("workflow"))),
        "",
        "### 证据与缺项规则",
        "",
    ]
    for rule in _string_list(blueprint.get("evidenceRules")):
        lines.append(f"- {rule}")
    for rule in _string_list(blueprint.get("missingRules")):
        lines.append(f"- {rule}")
    lines.extend([
        "- 同一问题只在最直接对应的评分项扣分一次，除非课程规则明确要求跨项影响。",
        "- 输出前核对总分等于分项合计，逐项理由均引用学生作业中的可观察证据。",
        "",
        "## 1. 评分项（score）",
        "",
        f"- `score_type`: `{score_type}`",
        f"- `full_score`: `{full_score}`",
    ])

    if score_type == "item":
        item_split = str(blueprint.get("itemSplit") or "按学生提交成果逐项拆分").strip()
        lines.extend([
            f"- `item_split`: {item_split}",
            f"- `item_count`: `{len(items)}`",
            "",
            "| 编号 | 分值 | 评分项内容 | 评分描述 | 评分规则 |",
            "|---:|---:|---|---|---|",
        ])
        for index, item in enumerate(items, start=1):
            rules = "；".join(_string_list(item.get("rules")))
            lines.append(
                f"| {index} | {int(item['score'])} | {_escape_table_cell(item.get('name'))} | "
                f"{_escape_table_cell(item.get('description'))} | {_escape_table_cell(rules)} |"
            )
        score_example = {
            "score": {
                "score_type": "item",
                "full_score": full_score,
                "total": 0,
                "items": [
                    {
                        "index": index,
                        "name": item.get("name"),
                        "score": 0,
                        "full": int(item["score"]),
                        "reason": "引用学生作业中的可观察证据",
                    }
                    for index, item in enumerate(items, start=1)
                ],
            }
        }
    else:
        lines.extend([
            "",
            "| 维度键 | 分值 | 维度名称 | 评分描述 | 评分规则 |",
            "|---|---:|---|---|---|",
        ])
        for index, item in enumerate(items, start=1):
            key = re.sub(r"[^a-z0-9_-]+", "-", str(item.get("key") or f"dimension-{index}").lower()).strip("-")
            rules = "；".join(_string_list(item.get("rules")))
            lines.append(
                f"| {key or f'dimension-{index}'} | {int(item['score'])} | {_escape_table_cell(item.get('name'))} | "
                f"{_escape_table_cell(item.get('description'))} | {_escape_table_cell(rules)} |"
            )
        score_example = {
            "score": {
                "score_type": "dimension",
                "full_score": full_score,
                "total": 0,
                "dimensions": [
                    {
                        "key": re.sub(r"[^a-z0-9_-]+", "-", str(item.get("key") or f"dimension-{index}").lower()).strip("-"),
                        "name": item.get("name"),
                        "score": 0,
                        "full": int(item["score"]),
                        "reason": "引用学生作业中的可观察证据",
                    }
                    for index, item in enumerate(items, start=1)
                ],
            }
        }

    evaluation_items = [item for item in blueprint.get("evaluationItems", []) if isinstance(item, dict)]
    lines.extend([
        "",
        "固定输出示例：",
        "",
        "```json",
        json.dumps(score_example, ensure_ascii=False, indent=2),
        "```",
        "",
        "## 2. 评价项（evaluations）",
        "",
        "| 评价项 | 生成要求 |",
        "|---|---|",
    ])
    for item in evaluation_items:
        lines.append(
            f"| {_escape_table_cell(item.get('name'))} | {_escape_table_cell(item.get('description'))} |"
        )
    lines.extend([
        "",
        "```json",
        json.dumps({
            "evaluations": {
                "综合评语": "基于分项证据概括整体表现",
                "改进建议": [
                    {"问题": "具体问题", "怎么改": "可执行修改动作", "目标": "修改后的质量标准"}
                ],
            }
        }, ensure_ascii=False, indent=2),
        "```",
        "",
        "输出前再次审计：分项合计正确、证据与结论一致、无重复扣分、评语与得分档位一致。",
        "",
    ])
    return "\n".join(lines)


def render_course_rules_reference(blueprint: Dict[str, Any]) -> str:
    lines = [
        "# 课程规则与校准说明",
        "",
        "## 课程硬规则",
        "",
    ]
    course_rules = _string_list(blueprint.get("courseRules")) or ["以学生提交要求和评分项为准。"]
    lines.extend(f"- {item}" for item in course_rules)
    lines.extend(["", "## 匿名校准说明", ""])
    calibration = _string_list(blueprint.get("calibrationNotes")) or ["暂无教师样本校准，按权威量表直接评分。"]
    lines.extend(f"- {item}" for item in calibration)
    lines.extend([
        "",
        "## 执行约束",
        "",
        "- 不读取文件路径中的分数、等级、姓名或学号作为评分依据。",
        "- 不复制历史学生原文；只使用匿名、可观察的质量特征校准尺度。",
        "- 未被教师样本覆盖的模块仍按课程权威量表评分。",
        "",
    ])
    return "\n".join(lines)


def build_grading_skill_zip(
    *,
    blueprint: Dict[str, Any],
    output_dir: Path,
) -> Dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    skill_name = str(blueprint["skillName"]).strip()
    skill_dir = output_dir / skill_name
    references_dir = skill_dir / "references"
    references_dir.mkdir(parents=True, exist_ok=True)
    (skill_dir / "README.md").write_text(README_TEMPLATE, encoding="utf-8")
    (skill_dir / "SKILL.md").write_text(render_skill_markdown(blueprint), encoding="utf-8")
    (references_dir / "course-rules.md").write_text(
        render_course_rules_reference(blueprint),
        encoding="utf-8",
    )

    zip_path = output_dir / f"{skill_name}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(skill_dir.rglob("*")):
            if path.is_file():
                archive.write(path, Path(skill_name) / path.relative_to(skill_dir))
    package_info = validate_grading_skill_package(str(zip_path))
    return {
        "zipPath": str(zip_path),
        "zipFileName": zip_path.name,
        "skillName": skill_name,
        "displayName": str(blueprint["displayName"]).strip(),
        "description": str(blueprint["description"]).strip(),
        "submissionRequirement": str(blueprint["submissionRequirement"]).strip(),
        "package": package_info,
    }


def generate_grading_skill_zip(
    *,
    material_paths: Sequence[Path],
    material_text: str,
    output_dir: Path,
    api_key: str,
    api_url: str,
    model: str,
) -> Dict[str, Any]:
    material_context = build_material_context(material_paths, material_text)
    blueprint = generate_skill_blueprint(
        material_context=material_context,
        api_key=api_key,
        api_url=api_url,
        model=model,
    )
    result = build_grading_skill_zip(blueprint=blueprint, output_dir=output_dir)
    result["blueprint"] = blueprint
    return result


def encode_file_base64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("ascii")


def select_sample_levels(count: int) -> List[tuple[str, str]]:
    normalized = min(5, max(1, int(count)))
    indexes_by_count = {
        1: [2],
        2: [0, 4],
        3: [0, 2, 4],
        4: [0, 1, 3, 4],
        5: [0, 1, 2, 3, 4],
    }
    return [SAMPLE_LEVELS[index] for index in indexes_by_count[normalized]]


def build_student_samples_prompt(
    *,
    assignment_title: str,
    submission_requirement: str,
    levels: Sequence[tuple[str, str]],
    validation_feedback: str = "",
) -> str:
    level_text = "\n".join(
        f"- 样例 {index}：内部目标档位为{level}；质量特征：{description}"
        for index, (level, description) in enumerate(levels, start=1)
    )
    feedback_section = (
        f"\n\n【上一次结果的校验问题】\n{validation_feedback}\n请修正全部问题。"
        if validation_feedback
        else ""
    )
    return f"""请为以下作业要求生成 {len(levels)} 份彼此独立的简单学生作业测试样例。

【作业标题】
{assignment_title or '课程作业'}

【学生提交作业要求】
{submission_requirement}

【内部测试档位】
{level_text}

要求：
1. 每份都像真实学生完成的 DOCX 正文，内容主题一致但表达、例子和完成质量不同。
2. 只生成文字、简单数据和可用纯文本表达的结构；不要声称已附上实际图片、视频或外部附件。
3. 不在标题、正文或文件名建议中出现“优秀、良好、中等、合格、较差、AI生成、测试样例、分数”等标签。
4. 低档样例要自然体现遗漏、浅层分析或少量错误，不能用元评论说明自己故意写差。
5. 每份包含 2 至 6 个章节，每章 1 至 4 个段落；禁止 LaTeX 和 Markdown 表格。
6. 只输出合法 JSON 对象，不要代码块或解释。

JSON 结构：
{{
  "assignments": [
    {{
      "title": "中性作业标题",
      "sections": [
        {{"heading": "章节标题", "paragraphs": ["完整自然段", "完整自然段"]}}
      ]
    }}
  ]
}}

assignments 数量必须恰好为 {len(levels)}，顺序与内部测试档位顺序一致。{feedback_section}"""


def _normalize_student_sample_blueprints(
    payload: Dict[str, Any],
    *,
    assignment_title: str,
    expected_count: int,
) -> List[Dict[str, Any]]:
    assignments = payload.get("assignments")
    if not isinstance(assignments, list) or len(assignments) != expected_count:
        raise SkillGenerationError(f"大模型应返回 {expected_count} 份学生作业")

    normalized: List[Dict[str, Any]] = []
    for index, item in enumerate(assignments, start=1):
        if not isinstance(item, dict):
            raise SkillGenerationError(f"第 {index} 份学生作业格式错误")
        title = str(item.get("title") or assignment_title or "课程作业").strip()
        sections = item.get("sections")
        if not isinstance(sections, list) or not sections:
            raise SkillGenerationError(f"第 {index} 份学生作业缺少章节")
        valid_sections = []
        for section in sections:
            if not isinstance(section, dict):
                continue
            heading = str(section.get("heading") or "正文").strip()
            paragraphs = _string_list(section.get("paragraphs"))
            if paragraphs:
                valid_sections.append({"heading": heading, "paragraphs": paragraphs})
        if not valid_sections:
            raise SkillGenerationError(f"第 {index} 份学生作业正文为空")

        visible_text = "\n".join(
            [title]
            + [section["heading"] for section in valid_sections]
            + [
                paragraph
                for section in valid_sections
                for paragraph in section["paragraphs"]
            ]
        )
        leaked = [token for token in SAMPLE_LABEL_LEAK_PATTERNS if token in visible_text]
        if leaked:
            raise SkillGenerationError(
                f"第 {index} 份学生作业泄漏了内部测试标签：{'/'.join(leaked)}"
            )
        normalized.append({"title": title, "sections": valid_sections})
    return normalized


def generate_student_sample_blueprints(
    *,
    assignment_title: str,
    submission_requirement: str,
    levels: Sequence[tuple[str, str]],
    api_key: str,
    api_url: str,
    model: str,
    max_attempts: int = 2,
) -> List[Dict[str, Any]]:
    feedback = ""
    last_error = ""
    for _ in range(max(1, max_attempts)):
        response_text = call_agenteval_llm(
            system_prompt=(
                "你是课程测试数据生成助手。生成内容不同质量档位的匿名学生作业，"
                "不得在作业正文泄漏档位标签，只输出 JSON。"
            ),
            user_prompt=build_student_samples_prompt(
                assignment_title=assignment_title,
                submission_requirement=submission_requirement,
                levels=levels,
                validation_feedback=feedback,
            ),
            api_key=api_key,
            api_url=api_url,
            model=model,
            max_tokens=10_000,
        )
        try:
            payload = parse_json_object(response_text)
            return _normalize_student_sample_blueprints(
                payload,
                assignment_title=assignment_title,
                expected_count=len(levels),
            )
        except SkillGenerationError as exc:
            last_error = str(exc)
            feedback = last_error
    raise SkillGenerationError(f"大模型生成的学生作业未通过校验：{last_error}")


def _set_run_font(run: Any, *, latin: str = "PingFang SC", east_asia: str = "PingFang SC", size: int = 11) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = latin
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")


def create_student_sample_docx(sample: Dict[str, Any], output_path: Path) -> None:
    """使用 compact_reference_guide 风格生成简单、稳定的学生 DOCX。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "PingFang SC")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "PingFang SC"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "PingFang SC")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Student Sample Title" not in [style.name for style in document.styles]:
        title_style = document.styles.add_style("Student Sample Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = document.styles["Student Sample Title"]
    title_style.font.name = "PingFang SC"
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(31, 77, 120)
    title_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title_style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "PingFang SC")
    title_style._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(14)
    title_style.paragraph_format.keep_with_next = True

    title_paragraph = document.add_paragraph(style="Student Sample Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(str(sample.get("title") or "课程作业"))
    _set_run_font(title_run, size=20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 77, 120)

    for section_data in sample.get("sections", []):
        document.add_heading(str(section_data.get("heading") or "正文"), level=1)
        for paragraph_text in section_data.get("paragraphs", []):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            run = paragraph.add_run(str(paragraph_text).strip())
            _set_run_font(run, size=11)

    document.core_properties.title = str(sample.get("title") or "课程作业")
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.save(str(output_path))


def generate_student_sample_docx_files(
    *,
    assignment_title: str,
    submission_requirement: str,
    count: int,
    output_dir: Path,
    api_key: str,
    api_url: str,
    model: str,
) -> List[Dict[str, Any]]:
    if len(submission_requirement.strip()) < 20:
        raise SkillGenerationError("请先生成或填写完整的学生作业要求")
    levels = select_sample_levels(count)
    blueprints = generate_student_sample_blueprints(
        assignment_title=assignment_title,
        submission_requirement=submission_requirement,
        levels=levels,
        api_key=api_key,
        api_url=api_url,
        model=model,
    )
    output_dir.mkdir(parents=True, exist_ok=True)
    results = []
    for index, (sample, level_info) in enumerate(zip(blueprints, levels), start=1):
        path = output_dir / f"学生作业_{index:02d}.docx"
        create_student_sample_docx(sample, path)
        results.append({
            "name": path.name,
            "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "base64": encode_file_base64(path),
            "level": level_info[0],
            "size": path.stat().st_size,
        })
    return results


def ensure_generation_error(exc: Exception) -> SkillGenerationError:
    if isinstance(exc, SkillGenerationError):
        return exc
    if isinstance(exc, CorrectionSkillError):
        return SkillGenerationError(str(exc))
    return SkillGenerationError(str(exc))
